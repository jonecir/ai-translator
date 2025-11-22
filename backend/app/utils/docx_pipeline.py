# backend/app/utils/docx_pipeline.py
from docx import Document
from app.utils.translator import translate_text  # ✅ import absoluto


def run_docx_to_docx(in_path: str, out_path: str, glossary: dict[str, str], source_lang="pt-BR", target_lang="en-US"):
    """
    Lê DOCX, traduz parágrafos e aplica substituições de glossário
    (após tradução). Salva em out_path. Retorna métricas simples.
    """
    doc = Document(in_path)
    paras = [p.text or "" for p in doc.paragraphs]

    # Tradução em lotes (lote de 50 para evitar payloads grandes)
    translated = []
    BATCH = 50
    for i in range(0, len(paras), BATCH):
        chunk = paras[i : i + BATCH]
        translated.extend(translate_text(chunk, source_lang, target_lang))

    # aplica glossário no texto traduzido
    def apply_gloss(txt: str) -> str:
        if not glossary:
            return txt
        out = txt
        for k, v in glossary.items():
            out = out.replace(k, v)
        return out

    # escreve de volta (perde formatação inline detalhada, mantém blocos)
    for p, new_text in zip(doc.paragraphs, translated):
        p.text = apply_gloss(new_text)

    doc.save(out_path)

    return {
        "paragraphs": len(paras),
        "source_lang": source_lang,
        "target_lang": target_lang,
        "glossary_terms": len(glossary or {}),
    }
