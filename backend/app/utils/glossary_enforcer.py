from collections import defaultdict
from docx import Document


def enforce_glossary(input_docx: str, output_docx: str, glossary_map: dict[str, str]) -> dict:
    """Simple DOCX→DOCX pass: replaces occurrences from glossary_map (src→dst),
    case-insensitive, returns basic metrics.
    """
    doc = Document(input_docx)
    replacements = defaultdict(int)
    def replace_in_run(run_text: str) -> str:
        new_text = run_text
        for src, dst in glossary_map.items():
            if not src:
                continue
            # naive case-insensitive replace preserving basic case
            new_text_lower = new_text.lower()
            src_lower = src.lower()
            if src_lower in new_text_lower:
                count = new_text_lower.count(src_lower)
                replacements[src] += count
                new_text = new_text.replace(src, dst)
                new_text = new_text.replace(src.capitalize(), dst)
                new_text = new_text.replace(src.upper(), dst.upper())
        return new_text


    for para in doc.paragraphs:
        for run in para.runs:
            run.text = replace_in_run(run.text)


    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.text = replace_in_run(run.text)


    doc.save(output_docx)

    metrics = {"replacements": dict(replacements), "unique_terms": len(replacements)}
    return metrics